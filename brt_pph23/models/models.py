# -*- coding: utf-8 -*-

from odoo import models, fields, api,exceptions
from odoo.exceptions import UserError, AccessError, ValidationError
from openerp.exceptions import except_orm, Warning, RedirectWarning
from datetime import date, datetime, timedelta

from datetime import date
from os.path import join as pjoin
import pdfkit
import shutil
from docx import Document
from docxtpl import DocxTemplate
import re
import docx

# pip install pdfkit
# pip install pdfcrowd
# pip install wkhtmltopdf
# pip install --pre python-docx
# pip install docxtpl

class account_invoice(models.Model):
    _inherit    = 'account.invoice'


    @api.multi
    def print_pph23(self):
        npwp        = str(self.partner_id.npwp)
        nama        = str(self.partner_id.name)
        alamat      = str(self.partner_id.street)
        id_inv      = self.id

        self.env.cr.execute("""
                            SELECT sum(s.price_subtotal) as tot_pajak, a.tax_id
                            FROM  account_invoice_line s, account_invoice_line_tax a
                            WHERE s.id = a.invoice_line_id
                            AND s.invoice_id = %s
                            GROUP BY a.tax_id
                            """,(self.id,))
        get_inv_line  = self.env.cr.dictfetchall()
        

        print('================================================================ data Call',get_inv_line)
        # print('================================================================ price_subtotal',price_subtotal)
        today       = date.today()
        bulan       = today.strftime("%B")
        thn         = today.strftime("%Y")
        tgl         = today.strftime("%d")
        
        bln_romawi  = ['I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX', 'X', 'XI', 'XII']
        get_bln_now = int(datetime.now().strftime("%m"))   

        doc_master  = "/opt/odoo/addons_custom/brt_pph23/static/src/doc/PPH23.docx"
        doc_copy    = "/opt/odoo/addons_custom/brt_pph23/static/src/doc/PPH23_salinan.docx"
        # doc_copy = "Free_Tracking_Manager_US.docx."+username+".docx"
        shutil.copy(doc_master, doc_copy)

        document    = docx.Document(doc_master)
        filename    = "/opt/odoo/addons_custom/brt_pph23/static/src/doc/PPH23_salinan.docx"

        doc         = Document(filename)
        tables      = [table for table in doc.tables];
        
        for data in get_inv_line :
            if data['tax_id'] == 4 :
                for table in tables:
                    for row in table.rows:
                        for cell in row.cells:
                            tot_pajak       = str(data['tot_pajak'])
                            tot_pajak       = str(data['tot_pajak'])
                            persen          = 15
                            sratus          = 100
                            potongan        = float(tot_pajak) * float(persen) / float(sratus)
                            cell.text       = cell.text.replace("[NILAI_DVD]",str(potongan))
                            cell.text       = cell.text.replace("[DEVIDEN]",str(tot_pajak))

            if data['tax_id'] == 5 :
                for table in tables:
                    for row in table.rows:
                        for cell in row.cells:
                            tot_pajak       = str(data['tot_pajak'])
                            tot_pajak       = str(data['tot_pajak'])
                            persen          = 15
                            sratus          = 100
                            potongan        = float(tot_pajak) * float(persen) / float(sratus)
                            cell.text       = cell.text.replace("[NILAI_BNG]",str(potongan))
                            cell.text       = cell.text.replace("[BUNGA]",str(tot_pajak))

            if data['tax_id'] == 6 :
                for table in tables:
                    for row in table.rows:
                        for cell in row.cells:
                            tot_pajak       = str(data['tot_pajak'])
                            tot_pajak       = str(data['tot_pajak'])
                            persen          = 15
                            sratus          = 100
                            potongan        = float(tot_pajak) * float(persen) / float(sratus)
                            cell.text       = cell.text.replace("[NILAI_RYT]",str(potongan))
                            cell.text       = cell.text.replace("[ROYALTI]",str(tot_pajak))

            if data['tax_id'] == 7 :
                for table in tables:
                    for row in table.rows:
                        for cell in row.cells:
                            tot_pajak       = str(data['tot_pajak'])
                            tot_pajak       = str(data['tot_pajak'])
                            persen          = 15
                            sratus          = 100
                            potongan        = float(tot_pajak) * float(persen) / float(sratus)
                            cell.text       = cell.text.replace("[NILAI_HRG]",str(potongan))
                            cell.text       = cell.text.replace("[PENGHARGAAN]",str(tot_pajak))

            if data['tax_id'] == 8 :
                for table in tables:
                    for row in table.rows:
                        for cell in row.cells:
                            tot_pajak       = str(data['tot_pajak'])
                            tot_pajak       = str(data['tot_pajak'])
                            persen          = 2
                            sratus          = 100
                            potongan        = float(tot_pajak) * float(persen) / float(sratus)
                            cell.text       = cell.text.replace("[NILAI_HRT]",str(potongan))
                            cell.text       = cell.text.replace("[SW_HARTA]",str(tot_pajak))
            
            if data['tax_id'] == 9 :
                for table in tables:
                    for row in table.rows:
                        for cell in row.cells:
                            tot_pajak       = str(data['tot_pajak'])
                            tot_pajak       = str(data['tot_pajak'])
                            persen          = 2
                            sratus          = 100
                            potongan        = float(tot_pajak) * float(persen) / float(sratus)
                            cell.text       = cell.text.replace("[NILAI_TKN]",str(potongan))
                            cell.text       = cell.text.replace("[JS_TEKNIK]",str(tot_pajak))

            if data['tax_id'] == 10 :
                for table in tables:
                    for row in table.rows:
                        for cell in row.cells:
                            tot_pajak       = str(data['tot_pajak'])
                            tot_pajak       = str(data['tot_pajak'])
                            persen          = 2
                            sratus          = 100
                            potongan        = float(tot_pajak) * float(persen) / float(sratus)
                            cell.text       = cell.text.replace("[NILAI_MGN]",str(potongan))
                            cell.text       = cell.text.replace("[JS_MANAGEMENT]",str(tot_pajak))

            if data['tax_id'] == 11 :
                for table in tables:
                    for row in table.rows:
                        for cell in row.cells:
                            tot_pajak       = str(data['tot_pajak'])
                            tot_pajak       = str(data['tot_pajak'])
                            persen          = 2
                            sratus          = 100
                            potongan        = float(tot_pajak) * float(persen) / float(sratus)
                            cell.text       = cell.text.replace("[NILAI_KST]",str(potongan))
                            cell.text       = cell.text.replace("[JS_KONSULTAN]",str(tot_pajak))
            
            if data['tax_id'] == 12 :
                for table in tables:
                    for row in table.rows:
                        for cell in row.cells:
                            tot_pajak       = str(data['tot_pajak'])
                            tot_pajak       = str(data['tot_pajak'])
                            persen          = 2
                            sratus          = 100
                            potongan        = float(tot_pajak) * float(persen) / float(sratus)
                            cell.text       = cell.text.replace("[NILAI_KST]",str(potongan))
                            cell.text       = cell.text.replace("[JS_KONSULTAN]",str(tot_pajak))

            if data['tax_id'] == 13 :
                for table in tables:
                    for row in table.rows:
                        for cell in row.cells:
                            tot_pajak       = str(data['tot_pajak'])
                            tot_pajak       = str(data['tot_pajak'])
                            persen          = 2
                            sratus          = 100
                            potongan        = float(tot_pajak) * float(persen) / float(sratus)
                            cell.text       = cell.text.replace("[NILAI_JSL_1]",str(potongan))
                            cell.text       = cell.text.replace("[JSL_1]",str(tot_pajak))   

            if data['tax_id'] == 14 :
                for table in tables:
                    for row in table.rows:
                        for cell in row.cells:
                            tot_pajak       = str(data['tot_pajak'])
                            tot_pajak       = str(data['tot_pajak'])
                            persen          = 2
                            sratus          = 100
                            potongan        = float(tot_pajak) * float(persen) / float(sratus)
                            cell.text       = cell.text.replace("[NILAI_JSL_2]",str(potongan))
                            cell.text       = cell.text.replace("[JSL_2]",str(tot_pajak))   

            if data['tax_id'] == 15 :
                for table in tables:
                    for row in table.rows:
                        for cell in row.cells:
                            tot_pajak       = str(data['tot_pajak'])
                            tot_pajak       = str(data['tot_pajak'])
                            persen          = 2
                            sratus          = 100
                            potongan        = float(tot_pajak) * float(persen) / float(sratus)
                            cell.text       = cell.text.replace("[NILAI_JSL_3]",str(potongan))
                            cell.text       = cell.text.replace("[JSL_3]",str(tot_pajak))   

            if data['tax_id'] == 16 :
                for table in tables:
                    for row in table.rows:
                        for cell in row.cells:
                            tot_pajak       = str(data['tot_pajak'])
                            tot_pajak       = str(data['tot_pajak'])
                            persen          = 2
                            sratus          = 100
                            potongan        = float(tot_pajak) * float(persen) / float(sratus)
                            cell.text       = cell.text.replace("[NILAI_JSL_4]",str(potongan))
                            cell.text       = cell.text.replace("[JSL_4]",str(tot_pajak))  

            if data['tax_id'] == 17 :
                for table in tables:
                    for row in table.rows:
                        for cell in row.cells:
                            tot_pajak       = str(data['tot_pajak'])
                            tot_pajak       = str(data['tot_pajak'])
                            persen          = 2
                            sratus          = 100
                            potongan        = float(tot_pajak) * float(persen) / float(sratus)
                            cell.text       = cell.text.replace("[NILAI_JSL_5]",str(potongan))
                            cell.text       = cell.text.replace("[JSL_5]",str(tot_pajak))   

            if data['tax_id'] == 18 :
                for table in tables:
                    for row in table.rows:
                        for cell in row.cells:
                            tot_pajak       = str(data['tot_pajak'])
                            tot_pajak       = str(data['tot_pajak'])
                            persen          = 2
                            sratus          = 100
                            potongan        = float(tot_pajak) * float(persen) / float(sratus)
                            cell.text       = cell.text.replace("[NILAI_JSL_6]",str(potongan))
                            cell.text       = cell.text.replace("[JSL_6]",str(tot_pajak))                      
            # else :
            #     for table in tables:
            #         for row in table.rows:
            #             for cell in row.cells:
            #                 cell.text = cell.text.replace("[DEVIDEN]","0")
            #                 cell.text = cell.text.replace("[NILAI_DVD]","0")

                    # print('================================================================ tot_pajak',tot_pajak)
                    # print('================================================================ potongan',potongan)


            # if data['tax_id'] == 5 :
            #     for table in tables:
            #         for row in table.rows:
            #             for cell in row.cells:
            #                 tot_pajak       = int(data['tot_pajak'])
            #                 potongan        = int(tot_pajak * 15 / 100)
            #                 total_pjk_dvd   = tot_pajak - potongan
            #                 cell.text       = cell.text.replace("[BUNGA]",tot_pajak)
            #                 cell.text       = cell.text.replace("[NILAI_BNG]",total_pjk_dvd)
            # else :
            #     for table in tables:
            #         for row in table.rows:
            #             for cell in row.cells:
            #                 cell.text = cell.text.replace("[BUNGA]","0")
            #                 cell.text = cell.text.replace("[NILAI_BNG]","0")


        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    if "A1" in cell.text:
                        cell.text = cell.text.replace("A1",npwp[0:1])
                    if "A2" in cell.text:
                        cell.text = cell.text.replace("A2",npwp[1:2])
                    if "A3" in cell.text:
                        cell.text = cell.text.replace("A3",npwp[2:3])
                    if "A4" in cell.text:
                        cell.text = cell.text.replace("A4",npwp[3:4])
                    if "A5" in cell.text:
                        cell.text = cell.text.replace("A5",npwp[4:5])
                    if "A6" in cell.text:
                        cell.text = cell.text.replace("A6",npwp[5:6])
                    if "A7" in cell.text:
                        cell.text = cell.text.replace("A7",npwp[6:7])
                    if "A8" in cell.text:
                        cell.text = cell.text.replace("A8",npwp[7:8])
                    if "A9" in cell.text:
                        cell.text = cell.text.replace("A9",npwp[8:9])

                    if "B1" in cell.text:
                        cell.text = cell.text.replace("B1",npwp[9:10])
                    if "B2" in cell.text:
                        cell.text = cell.text.replace("B2",npwp[10:11])
                    if "B3" in cell.text:
                        cell.text = cell.text.replace("B3",npwp[11:12])
                    if "B4" in cell.text:
                        cell.text = cell.text.replace("B4",npwp[12:13])
                    if "B5" in cell.text:
                        cell.text = cell.text.replace("B5",npwp[13:14])
                    if "B6" in cell.text:
                        cell.text = cell.text.replace("B6",npwp[14:15])
                    if "B7" in cell.text:
                        cell.text = cell.text.replace("B7",npwp[15:16])
                    if "B8" in cell.text:
                        cell.text = cell.text.replace("B8",npwp[16:17])
                    if "B9" in cell.text:
                        cell.text = cell.text.replace("B9",npwp[17:18])

                    if "C1" in cell.text:
                        cell.text = cell.text.replace("C1",nama[0:1])
                    if "C2" in cell.text:
                        cell.text = cell.text.replace("C2",nama[1:2])
                    if "C3" in cell.text:
                        cell.text = cell.text.replace("C3",nama[2:3])
                    if "C4" in cell.text:
                        cell.text = cell.text.replace("C4",nama[3:4])
                    if "C5" in cell.text:
                        cell.text = cell.text.replace("C5",nama[4:5])
                    if "C6" in cell.text:
                        cell.text = cell.text.replace("C6",nama[5:6])
                    if "C7" in cell.text:
                        cell.text = cell.text.replace("C7",nama[6:7])
                    if "C8" in cell.text:
                        cell.text = cell.text.replace("C8",nama[7:8])
                    if "C9" in cell.text:
                        cell.text = cell.text.replace("C9",nama[8:9])

                    if "D1" in cell.text:
                        cell.text = cell.text.replace("D1",nama[9:10])
                    if "D2" in cell.text:
                        cell.text = cell.text.replace("D2",nama[10:11])
                    if "D3" in cell.text:
                        cell.text = cell.text.replace("D3",nama[11:12])
                    if "D4" in cell.text:
                        cell.text = cell.text.replace("D4",nama[12:13])
                    if "D5" in cell.text:
                        cell.text = cell.text.replace("D5",nama[13:14])
                    if "D6" in cell.text:
                        cell.text = cell.text.replace("D6",nama[14:15])
                    if "D7" in cell.text:
                        cell.text = cell.text.replace("D7",nama[15:16])
                    if "D8" in cell.text:
                        cell.text = cell.text.replace("D8",nama[16:17])
                    if "D9" in cell.text:
                        cell.text = cell.text.replace("D9",nama[17:18])

                    if "E1" in cell.text:
                        cell.text = cell.text.replace("E1",alamat[0:1])
                    if "E2" in cell.text:
                        cell.text = cell.text.replace("E2",alamat[1:2])
                    if "E3" in cell.text:
                        cell.text = cell.text.replace("E3",alamat[2:3])
                    if "E4" in cell.text:
                        cell.text = cell.text.replace("E4",alamat[3:4])
                    if "E5" in cell.text:
                        cell.text = cell.text.replace("E5",alamat[4:5])
                    if "E6" in cell.text:
                        cell.text = cell.text.replace("E6",alamat[5:6])
                    if "E7" in cell.text:
                        cell.text = cell.text.replace("E7",alamat[6:7])
                    if "E8" in cell.text:
                        cell.text = cell.text.replace("E8",alamat[7:8])
                    if "E9" in cell.text:
                        cell.text = cell.text.replace("E9",alamat[8:9])

                    if "F1" in cell.text:
                        cell.text = cell.text.replace("F1",alamat[9:10])
                    if "F2" in cell.text:
                        cell.text = cell.text.replace("F2",alamat[10:11])
                    if "F3" in cell.text:
                        cell.text = cell.text.replace("F3",alamat[11:12])
                    if "F4" in cell.text:
                        cell.text = cell.text.replace("F4",alamat[12:13])
                    if "F5" in cell.text:
                        cell.text = cell.text.replace("F5",alamat[13:14])
                    if "F6" in cell.text:
                        cell.text = cell.text.replace("F6",alamat[14:15])
                    if "F7" in cell.text:
                        cell.text = cell.text.replace("F7",alamat[15:16])
                    if "F8" in cell.text:
                        cell.text = cell.text.replace("F8",alamat[16:17])
                    if "F9" in cell.text:
                        cell.text = cell.text.replace("F9",alamat[17:18])

        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    if "[npwp]" in cell.text:
                        cell.text = cell.text.replace("[npwp]",npwp)
                        cell.text = cell.text.replace("[nmwp]",nama)
                        cell.text = cell.text.replace("[bulan]",bulan)
                        cell.text = cell.text.replace("[tgl]",tgl)
                        cell.text = cell.text.replace("[thn]",thn)
                        cell.text = cell.text.replace("[bln_romawi]",bln_romawi[get_bln_now - 1])

        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    cell.text = cell.text.replace("[DEVIDEN]","0")
                    cell.text = cell.text.replace("[BUNGA]","0")
                    cell.text = cell.text.replace("[ROYALTI]","0")
                    cell.text = cell.text.replace("[PENGHARGAAN]","0")
                    cell.text = cell.text.replace("[SW_HARTA]","0")
                    cell.text = cell.text.replace("[JS_TEKNIK]","0")
                    cell.text = cell.text.replace("[JS_MANAGEMENT]","0")
                    cell.text = cell.text.replace("[JS_KONSULTAN]","0")
                    cell.text = cell.text.replace("[JSL_1]","0")
                    cell.text = cell.text.replace("[JSL_2]","0")
                    cell.text = cell.text.replace("[JSL_3]","0")
                    cell.text = cell.text.replace("[JSL_4]","0")
                    cell.text = cell.text.replace("[JSL_5]","0")
                    cell.text = cell.text.replace("[JSL_6]","0")

                    cell.text = cell.text.replace("[NILAI_DVD]","0")
                    cell.text = cell.text.replace("[NILAI_BNG]","0")
                    cell.text = cell.text.replace("[NILAI_RYT]","0")
                    cell.text = cell.text.replace("[NILAI_HRG]","0")
                    cell.text = cell.text.replace("[NILAI_HRT]","0")
                    cell.text = cell.text.replace("[NILAI_TKN]","0")
                    cell.text = cell.text.replace("[NILAI_MGN]","0")
                    cell.text = cell.text.replace("[NILAI_KST]","0")
                    cell.text = cell.text.replace("[NILAI_JSL_1]","0")
                    cell.text = cell.text.replace("[NILAI_JSL_2]","0")
                    cell.text = cell.text.replace("[NILAI_JSL_3]","0")
                    cell.text = cell.text.replace("[NILAI_JSL_4]","0")
                    cell.text = cell.text.replace("[NILAI_JSL_5]","0")
                    cell.text = cell.text.replace("[NILAI_JSL_6]","0")

        # list1= ['ABC','XYZ']
        # list2 = ['BAYU','DEF']
        # for p in doc.paragraphs:
        #     inline = p.runs
        #     for j in range(0,len(inline)):
        #         for i in range(0, len(list1)):
        #             inline[j].text = inline[j].text.replace(list1[i], list2[i])
        #             print(p.text)
        #             # print(inline[j].text)
        doc.save('/opt/odoo/addons_custom/brt_pph23/static/src/doc/PPh23_cetak_'+nama+'.docx')
        
        return {
            'type'    : 'ir.actions.act_url',
            'url'     : '/brt_pph23/static/src/doc/PPh23_cetak_'+nama+'.docx',            
            # 'url'     : 'C:/odoo-10/8040-jaddi/addons_custom/brt_pph23/static/src/doc/ASD.docx',            
            # 'url': '/mcs_survey/static/'+str(self.title)+'.docx',            
            'target': 'new',
        }

      
        


    # @api.multi
    # def print_pph23(self):
    #     templatePPh23  		= self.env['brt_report.template_report'].search([('kode','=','RPP23')]) 
    #     str_templatePPh23   = templatePPh23.template
    #     space 		= " "
    #     npwp 		= str(self.partner_id.npwp)
    #     nmvendor 	= self.partner_id.name
        
    #     
    #     A10 = npwp[9:10]
    #     A11 = npwp[10:11]
    #     A12 = npwp[11:12]
    #     A13 = npwp[12:13]

    #     B1 = nmvendor[0:1]
    #     B2 = nmvendor[1:2]
    #     B3 = nmvendor[2:3]
    #     B4 = nmvendor[3:4]
    #     B5 = nmvendor[4:5]
    #     B6 = nmvendor[5:6]
    #     B7 = nmvendor[6:7]
    #     B8 = nmvendor[7:8]
    #     B9 = nmvendor[8:9]
    #     B10 = nmvendor[9:10]
    #     B11 = nmvendor[10:11]
    #     B12 = nmvendor[11:12]
    #     B13 = nmvendor[12:13]
    #     B14 = nmvendor[13:14]
    #     B15 = nmvendor[14:15]
    #     B16 = nmvendor[15:16]
    #     B17 = nmvendor[16:17]
    #     B18 = nmvendor[17:18]
    #     B19 = nmvendor[18:19]
    #     B20 = nmvendor[19:20]

    #     str_templatePPh23   = str_templatePPh23.replace("A10",A10)
    #     str_templatePPh23   = str_templatePPh23.replace("A11",A11)
    #     str_templatePPh23   = str_templatePPh23.replace("A12",A12)
    #     str_templatePPh23   = str_templatePPh23.replace("A13",A13)

    #     str_templatePPh23   = str_templatePPh23.replace("A1",A1)
    #     str_templatePPh23   = str_templatePPh23.replace("A2",A2)
    #     str_templatePPh23   = str_templatePPh23.replace("A3",A3)
    #     str_templatePPh23   = str_templatePPh23.replace("A4",A4)
    #     str_templatePPh23   = str_templatePPh23.replace("A5",A5)
    #     str_templatePPh23   = str_templatePPh23.replace("A6",A6)
    #     str_templatePPh23   = str_templatePPh23.replace("A7",A7)
    #     str_templatePPh23   = str_templatePPh23.replace("A8",A8)
    #     str_templatePPh23   = str_templatePPh23.replace("A9",A9)
        

    #     str_templatePPh23   = str_templatePPh23.replace("B20",space)
    #     str_templatePPh23   = str_templatePPh23.replace("B21",space)
    #     str_templatePPh23   = str_templatePPh23.replace("B22",space)
    #     str_templatePPh23   = str_templatePPh23.replace("B23",space)

    #     str_templatePPh23   = str_templatePPh23.replace("B10",B10)
    #     str_templatePPh23   = str_templatePPh23.replace("B11",B11)
    #     str_templatePPh23   = str_templatePPh23.replace("B12",B12)
    #     str_templatePPh23   = str_templatePPh23.replace("B13",B13)
    #     str_templatePPh23   = str_templatePPh23.replace("B14",B14)
    #     str_templatePPh23   = str_templatePPh23.replace("B15",B15)
    #     str_templatePPh23   = str_templatePPh23.replace("B16",B16)
    #     str_templatePPh23   = str_templatePPh23.replace("B17",B17)
    #     str_templatePPh23   = str_templatePPh23.replace("B18",B18)
    #     str_templatePPh23   = str_templatePPh23.replace("B19",B19)

    #     str_templatePPh23   = str_templatePPh23.replace("B1",B1)
    #     str_templatePPh23   = str_templatePPh23.replace("B2",B2)
    #     str_templatePPh23   = str_templatePPh23.replace("B3",B3)
    #     str_templatePPh23   = str_templatePPh23.replace("B4",B4)
    #     str_templatePPh23   = str_templatePPh23.replace("B5",B5)
    #     str_templatePPh23   = str_templatePPh23.replace("B6",B6)
    #     str_templatePPh23   = str_templatePPh23.replace("B7",B7)
    #     str_templatePPh23   = str_templatePPh23.replace("B8",B8)
    #     str_templatePPh23   = str_templatePPh23.replace("B9",B9)
        
    #     A1 = " "
    #     A2 = " "
    #     A3 = " "
    #     A4 = " "
    #     A5 = " "
    #     A6 = " "
    #     A7 = " "
    #     A8 = " "
    #     A9 = " "
    #     A10 = " "
    #     A11 = " "
    #     A12 = " "
    #     A13 = " "

    #     str_templatePPh23   = str_templatePPh23.replace("A10",A10)
    #     str_templatePPh23   = str_templatePPh23.replace("A11",A11)
    #     str_templatePPh23   = str_templatePPh23.replace("A12",A12)
    #     str_templatePPh23   = str_templatePPh23.replace("A13",A13)
        
    #     str_templatePPh23   = str_templatePPh23.replace("A1",A1)
    #     str_templatePPh23   = str_templatePPh23.replace("A2",A2)
    #     str_templatePPh23   = str_templatePPh23.replace("A3",A3)
    #     str_templatePPh23   = str_templatePPh23.replace("A4",A4)
    #     str_templatePPh23   = str_templatePPh23.replace("A5",A5)
    #     str_templatePPh23   = str_templatePPh23.replace("A6",A6)
    #     str_templatePPh23   = str_templatePPh23.replace("A7",A7)
    #     str_templatePPh23   = str_templatePPh23.replace("A8",A8)
    #     str_templatePPh23   = str_templatePPh23.replace("A9",A9)
        

    #     str_templatePPh23   = str_templatePPh23.replace("XX",str(space))

    # 	id_po 				= self.id
    #     config 				= pdfkit.configuration(wkhtmltopdf='C:/Program Files/wkhtmltopdf/bin/wkhtmltopdf.exe')
    #     pdfkit.from_string(str_templatePPh23, 'C:/odoo-10/8040-jaddi/addons_custom/brt_pph23/static/src/pdf/'+str(id_po)+'-pph23.pdf', configuration=config)
    #    	return {
    #         'type' 	: 'ir.actions.act_url',
    #         'url' 	: '/brt_pph23/static/src/pdf/'+str(id_po)+'-pph23.pdf',            
    #         # 'url': '/mcs_survey/static/'+str(self.title)+'.docx',            
    #         'target': 'new',
    #     }

   

    @api.multi
    def print_pph_total(self):
        npwp        = str(self.partner_id.npwp)
        nama        = str(self.partner_id.name)
        alamat      = str(self.partner_id.street)
        id_inv      = self.id

        self.env.cr.execute("""
                            SELECT sum(s.price_subtotal) as tot_pajak, a.tax_id
                            FROM  account_invoice_line s, account_invoice_line_tax a
                            WHERE s.id = a.invoice_line_id
                            AND s.invoice_id = %s
                            AND a.tax_id = 8
                            GROUP BY a.tax_id
                            """,(self.id,))
        get_inv_line  = self.env.cr.dictfetchall()
        

        print('================================================================ data Call',get_inv_line)
        # print('================================================================ price_subtotal',price_subtotal)
        today       = date.today()
        bulan       = today.strftime("%B")
        thn         = today.strftime("%Y")
        tgl         = today.strftime("%d")
        
        bln_romawi  = ['I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX', 'X', 'XI', 'XII']
        get_bln_now = int(datetime.now().strftime("%m"))   

        master_pphtotal  = "/opt/odoo/addons_custom/brt_pph23/static/src/doc/bukti_Potong.docx"
        pphtotal_copy    = "/opt/odoo/addons_custom/brt_pph23/static/src/doc/bukti_Potong_salinan.docx"
        shutil.copy(master_pphtotal, pphtotal_copy)

        document    = docx.Document(pphtotal_copy)
        filename    = "/opt/odoo/addons_custom/brt_pph23/static/src/doc/bukti_Potong_salinan.docx"

        doc         = Document(filename)
        tables      = [table for table in doc.tables];
        
        for data in get_inv_line :
            if data['tax_id'] == 8 :
                for table in tables:
                    for row in table.rows:
                        for cell in row.cells:
                            tot_pajak       = str(data['tot_pajak'])
                            tot_pajak       = str(data['tot_pajak'])
                            persen          = 10
                            sratus          = 100
                            potongan        = float(tot_pajak) * float(persen) / float(sratus)
                            cell.text       = cell.text.replace("[PPH_DIPOTONG]",str(potongan))
                            cell.text       = cell.text.replace("[JML_BRUTO]",str(tot_pajak))

                for table in tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if "A1" in cell.text:
                                cell.text = cell.text.replace("A1",npwp[0:1])
                            if "A2" in cell.text:
                                cell.text = cell.text.replace("A2",npwp[1:2])
                            if "A3" in cell.text:
                                cell.text = cell.text.replace("A3",npwp[2:3])
                            if "A4" in cell.text:
                                cell.text = cell.text.replace("A4",npwp[3:4])
                            if "A5" in cell.text:
                                cell.text = cell.text.replace("A5",npwp[4:5])
                            if "A6" in cell.text:
                                cell.text = cell.text.replace("A6",npwp[5:6])
                            if "A7" in cell.text:
                                cell.text = cell.text.replace("A7",npwp[6:7])
                            if "A8" in cell.text:
                                cell.text = cell.text.replace("A8",npwp[7:8])
                            if "A9" in cell.text:
                                cell.text = cell.text.replace("A9",npwp[8:9])

                            if "B1" in cell.text:
                                cell.text = cell.text.replace("B1",npwp[9:10])
                            if "B2" in cell.text:
                                cell.text = cell.text.replace("B2",npwp[10:11])
                            if "B3" in cell.text:
                                cell.text = cell.text.replace("B3",npwp[11:12])
                            if "B4" in cell.text:
                                cell.text = cell.text.replace("B4",npwp[12:13])
                            if "B5" in cell.text:
                                cell.text = cell.text.replace("B5",npwp[13:14])
                            if "B6" in cell.text:
                                cell.text = cell.text.replace("B6",npwp[14:15])

                            if "C1" in cell.text:
                                cell.text = cell.text.replace("C1",nama[0:1])
                            if "C2" in cell.text:
                                cell.text = cell.text.replace("C2",nama[1:2])
                            if "C3" in cell.text:
                                cell.text = cell.text.replace("C3",nama[2:3])
                            if "C4" in cell.text:
                                cell.text = cell.text.replace("C4",nama[3:4])
                            if "C5" in cell.text:
                                cell.text = cell.text.replace("C5",nama[4:5])
                            if "C6" in cell.text:
                                cell.text = cell.text.replace("C6",nama[5:6])
                            if "C7" in cell.text:
                                cell.text = cell.text.replace("C7",nama[6:7])
                            if "C8" in cell.text:
                                cell.text = cell.text.replace("C8",nama[7:8])
                            if "C9" in cell.text:
                                cell.text = cell.text.replace("C9",nama[8:9])

                            if "D1" in cell.text:
                                cell.text = cell.text.replace("D1",nama[9:10])
                            if "D2" in cell.text:
                                cell.text = cell.text.replace("D2",nama[10:11])
                            if "D3" in cell.text:
                                cell.text = cell.text.replace("D3",nama[11:12])
                            if "D4" in cell.text:
                                cell.text = cell.text.replace("D4",nama[12:13])
                            if "D5" in cell.text:
                                cell.text = cell.text.replace("D5",nama[13:14])
                            if "D6" in cell.text:
                                cell.text = cell.text.replace("D6",nama[14:15])
                            if "D7" in cell.text:
                                cell.text = cell.text.replace("D7",nama[15:16])
                            if "D8" in cell.text:
                                cell.text = cell.text.replace("D8",nama[16:17])
                            if "D9" in cell.text:
                                cell.text = cell.text.replace("D9",nama[17:18])

                            if "E1" in cell.text:
                                cell.text = cell.text.replace("E1",nama[18:19])
                            if "E2" in cell.text:
                                cell.text = cell.text.replace("E2",nama[19:20])

                            if "F1" in cell.text:
                                cell.text = cell.text.replace("F1",alamat[0:1])
                            if "F2" in cell.text:
                                cell.text = cell.text.replace("F2",alamat[1:2])
                            if "F3" in cell.text:
                                cell.text = cell.text.replace("F3",alamat[2:3])
                            if "F4" in cell.text:
                                cell.text = cell.text.replace("F4",alamat[3:4])
                            if "F5" in cell.text:
                                cell.text = cell.text.replace("F5",alamat[4:5])
                            if "F6" in cell.text:
                                cell.text = cell.text.replace("F6",alamat[5:6])
                            if "F7" in cell.text:
                                cell.text = cell.text.replace("F7",alamat[6:7])
                            if "F8" in cell.text:
                                cell.text = cell.text.replace("F8",alamat[7:8])
                            if "F9" in cell.text:
                                cell.text = cell.text.replace("F9",alamat[8:9])

                            if "G1" in cell.text:
                                cell.text = cell.text.replace("G1",nama[9:10])
                            if "G2" in cell.text:
                                cell.text = cell.text.replace("G2",nama[10:11])
                            if "G3" in cell.text:
                                cell.text = cell.text.replace("G3",nama[11:12])
                            if "G4" in cell.text:
                                cell.text = cell.text.replace("G4",nama[12:13])
                            if "G5" in cell.text:
                                cell.text = cell.text.replace("G5",nama[13:14])
                            if "G6" in cell.text:
                                cell.text = cell.text.replace("G6",nama[14:15])
                            if "G7" in cell.text:
                                cell.text = cell.text.replace("G7",nama[15:16])
                            if "G8" in cell.text:
                                cell.text = cell.text.replace("G8",nama[16:17])
                            if "G9" in cell.text:
                                cell.text = cell.text.replace("G9",nama[17:18])

                            if "H1" in cell.text:
                                cell.text = cell.text.replace("H1",nama[18:19])
                            if "H2" in cell.text:
                                cell.text = cell.text.replace("H2",nama[19:20])

                for table in tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if "[npwp]" in cell.text:
                                cell.text = cell.text.replace("[npwp]",npwp)
                                cell.text = cell.text.replace("[nmwp]",nama)
                                cell.text = cell.text.replace("[bulan]",bulan)
                                cell.text = cell.text.replace("[tgl]",tgl)
                                cell.text = cell.text.replace("[thn]",thn)
                                cell.text = cell.text.replace("[bln_romawi]",bln_romawi[get_bln_now - 1])

                doc.save('/opt/odoo/addons_custom/brt_pph23/static/src/doc/Cetak_bukti_Potong.docx')
                
                return {
                    'type'    : 'ir.actions.act_url',
                    'url'     : '/brt_pph23/static/src/doc/Cetak_bukti_Potong.docx',            
                    # 'url'     : 'C:/odoo-10/8040-jaddi/addons_custom/brt_pph23/static/src/doc/ASD.docx',            
                    # 'url': '/mcs_survey/static/'+str(self.title)+'.docx',            
                    'target': 'new',
                }
            else:
                raise exceptions.ValidationError('Akses Ditolak, Transaksi Ini tidak memiliki PPh23 yang harus dipotong')


    @api.multi
    def abc(self):
    	id_po 				= self.id
        config 				= pdfkit.configuration(wkhtmltopdf='C:/Program Files/wkhtmltopdf/bin/wkhtmltopdf.exe')
        pdfkit.from_string(html, 'C:/odoo-10/8040-jaddi/addons_custom/brt_pph23/static/src/pdf/'+str(id_po)+'-pphtotal.pdf', configuration=config)
       	return {
            'type' 	: 'ir.actions.act_url',
            'url' 	: '/brt_pph23/static/src/pdf/'+str(id_po)+'-pphtotal.pdf',            
            # 'url': '/mcs_survey/static/'+str(self.title)+'.docx',            
            'target': 'new',
        }