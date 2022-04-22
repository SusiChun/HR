# -*- coding: utf-8 -*-

from odoo import models, fields, api,exceptions
from odoo.exceptions import UserError, AccessError, ValidationError
from openerp.exceptions import except_orm, Warning, RedirectWarning
from datetime import date, datetime, timedelta

from datetime import date
from os.path import join as pjoin
import pdfkit
import shutil
from docx import Document
from docxtpl import DocxTemplate
import re
import docx

class HRcontract(models.Model):
    _inherit = 'hr.contract'
    
    no_urut         = fields.Integer(string="No. Urut")
    no_kontrak      = fields.Char(string="Contract Number")

    def buat_nomor(self):
        print '=========================== tessssssssssssssss'
        bln_romawi  = ['I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX', 'X', 'XI', 'XII']
        get_thn_now = str(datetime.now().strftime("%Y"))
        get_bln_now = int(datetime.now().strftime("%m"))   
        self.env.cr.execute("""
                            SELECT
                                MAX(s.no_urut) AS no_max 
                            FROM
                                hr_contract s 
                            """,)
        get_no_max  = self.env.cr.dictfetchall()
        print '=============================================================', get_no_max
        next_no_max = 1
        for data in get_no_max :
            if data['no_max'] :
                next_no_max = int(data['no_max']) + 1

        no_kontrak        = str(next_no_max) + "/MJT/" + bln_romawi[get_bln_now - 1] + "-" + get_thn_now  
        self.no_kontrak   = no_kontrak
        self.no_urut      = next_no_max


    @api.multi
    def cetak_kontrak(self):
        id_kontrak      = str(self.id)
        no_urut      = str(self.no_urut)
        id_employee     = str(self.employee_id.id)
        employee        = str(self.employee_id.name)
        ktp             = str(self.employee_id.nik)
        religion        = str(self.employee_id.religion)
        birthday        = str(self.employee_id.birthday)
        street          = str(self.employee_id.address)
        telp            = str(self.employee_id.mobile_phone)
        job_id          = str(self.job_id.name)
        type_id         = str(self.type_id.name)
        date_start      = str(self.date_start)
        date_end        = str(self.date_end)

        today           = date.today()
        bulan           = today.strftime("%B")
        thn             = today.strftime("%Y")
        tgl             = today.strftime("%d")
        
        bln_romawi      = ['I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX', 'X', 'XI', 'XII']
        get_bln_now     = int(datetime.now().strftime("%m"))   
        romawi          = bln_romawi[get_bln_now - 1]
        doc_master      = "/opt/odoo/addons_custom/brt_hirs_report/static/src/doc/employee_contract_2019.docx"
        # doc_copy      = "/opt/odoo/addons_custom/brt_hirs_report/static/src/doc/salinan_kontrak.docx"
        # doc_copy = "Free_Tracking_Manager_US.docx."+username+".docx"
        # shutil.copy(doc_master, doc_copy)

        document    = docx.Document(doc_master)
        filename    = "/opt/odoo/addons_custom/brt_hirs_report/static/src/doc/employee_contract_2019.docx"

        doc         = Document(filename)
        
        list1       = ['nourut','blnromawi','thn','tgl','bln','employee','noktp','ttl','almt','agama','telp','contractType','idmployee','dateStart','dateEnd','jobTitle','employee']
        list2       = [no_urut, romawi, thn, tgl, bulan,employee,ktp,birthday,street,religion,telp,type_id,id_employee,date_start,date_end,job_id,employee]
        for p in doc.paragraphs:
            inline = p.runs
            for j in range(0,len(inline)):
                for i in range(0, len(list1)):
                    inline[j].text = inline[j].text.replace(list1[i], list2[i])
                    # print(p.text)
                    # print(inline[j].text)

        doc.save('/opt/odoo/addons_custom/brt_hirs_report/static/src/doc/cetak_kontrak.docx')
        
        return {
            'type'      : 'ir.actions.act_url',
            'url'       : '/brt_hirs_report/static/src/doc/cetak_kontrak.docx',            
            'target'    : 'new',
        }